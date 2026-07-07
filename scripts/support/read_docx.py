#!/usr/bin/env python3
"""Inspect DOCX files and optionally extract text and tables."""

from __future__ import annotations

import argparse
from pathlib import Path
from typing import Any

from common import detect_dependency, ensure_dir, exit_if_missing, file_summary, print_report, resolve_mode, utc_timestamp, write_text


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Inspect a DOCX file and optionally extract content.")
    parser.add_argument("input", type=Path)
    parser.add_argument("--report", action="store_true")
    parser.add_argument("--execute", action="store_true")
    parser.add_argument("--out", type=Path, help="Output directory for extracted text.")
    parser.add_argument("--sample-paragraphs", type=int, default=10)
    return parser.parse_args()


def inspect_docx(path: Path, sample_paragraphs: int) -> dict[str, Any]:
    if not detect_dependency("docx"):
        return {
            "dependency_missing": "python-docx",
            "paragraph_count": 0,
            "table_count": 0,
            "sample_paragraphs": [],
            "sample_tables": [],
        }

    import docx

    document = docx.Document(str(path))
    sample_paragraphs_data = [p.text.strip() for p in document.paragraphs if p.text.strip()][:sample_paragraphs]
    sample_tables = []
    for table in document.tables[:3]:
        rows = []
        for row in table.rows[:5]:
            rows.append([cell.text.strip() for cell in row.cells])
        sample_tables.append(rows)
    return {
        "paragraph_count": len(document.paragraphs),
        "table_count": len(document.tables),
        "sample_paragraphs": sample_paragraphs_data,
        "sample_tables": sample_tables,
    }


def extract_docx(path: Path) -> tuple[str, str]:
    import docx

    document = docx.Document(str(path))
    text_parts = []
    markdown_parts = []
    for paragraph in document.paragraphs:
        if paragraph.text.strip():
            text_parts.append(paragraph.text.strip())
            markdown_parts.append(paragraph.text.strip())
    for table_index, table in enumerate(document.tables, start=1):
        markdown_parts.append(f"\n## Table {table_index}")
        for row in table.rows:
            cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
            markdown_parts.append(" | ".join(cells))
    return "\n".join(text_parts).strip() + "\n", "\n".join(markdown_parts).strip() + "\n"


def main() -> int:
    args = parse_args()
    mode = resolve_mode(args)
    exit_if_missing(args.input)
    if args.input.suffix.lower() != ".docx":
        raise SystemExit("Unsupported file type. Use a .docx input.")
    info = inspect_docx(args.input, args.sample_paragraphs)
    report = {
        "mode": mode,
        "session_timestamp": utc_timestamp(),
        "input": file_summary(args.input),
        "inspection": info,
        "will_write": ["extracted.txt", "summary.md"] if mode == "execute" and "dependency_missing" not in info else [],
    }
    print_report(report)

    if mode != "execute":
        return 0
    if info.get("dependency_missing"):
        raise SystemExit("Cannot execute DOCX extraction: missing optional dependency python-docx.")

    text, markdown = extract_docx(args.input)
    out_dir = args.out or args.input.parent / f"{args.input.stem}-docx"
    ensure_dir(out_dir)
    write_text(out_dir / "extracted.txt", text)
    write_text(out_dir / "summary.md", f"Session timestamp: {utc_timestamp()}\n\n{markdown}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
